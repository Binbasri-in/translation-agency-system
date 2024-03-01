import fitz  # PyMuPDF
from docx import Document as DocxDocument
import arabic_reshaper
from bidi.algorithm import get_display
import pytesseract
from PIL import Image
import pandas as pd
import re
import requests
from ocr_project.settings import CLOUDFLARE_API_TOKEN, CLOUDFLARE_ACCOUNT_ID


def perform_ocr_image_to_text(file):
    print("Performing OCR on the image...")
    image = Image.open(file)
    text = pytesseract.image_to_string(image)
    output_file = "output.txt"
    with open(output_file, 'w') as f:
        f.write(text)
    return output_file


def perform_ocr_image_to_docx(file):
    print("Performing OCR on the image...")
    image = Image.open(file)
    text = pytesseract.image_to_string(image)
    # Filter out control characters and other non-XML-compatible characters
    filtered_text = re.sub(r'[^\x20-\x7E\x0A\x0D]', '', text)
    docx = DocxDocument()
    docx.add_paragraph(filtered_text)
    output_file = "output.docx"
    docx.save(output_file)
    return output_file


def perform_ocr_image_to_xlsx(file):
    print("Performing OCR on the image...")
    image = Image.open(file)
    text = pytesseract.image_to_string(image)
    # Filter out illegal characters for Excel
    filtered_text = re.sub(r'[^\x20-\x7E\x0A\x0D]', '', text)
    df = pd.DataFrame([filtered_text.split('\n')])
    output_file = "output.xlsx"
    df.to_excel(output_file, index=False, header=False)
    return output_file



def perform_ocr_pdf_to_docx(file):
    print("Performing OCR on each page in the PDF file...")
    pdf_document = fitz.open(stream=file.read(), filetype="pdf")
    docx = DocxDocument()  # Create a new Word document
    for page_number in range(len(pdf_document)):
        page = pdf_document.load_page(page_number)
        pix = page.get_pixmap()
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        text = pytesseract.image_to_string(img)
        # Reshape and reorder Arabic text
        reshaped_text = arabic_reshaper.reshape(text)
        for line in reshaped_text.split('\n'):
            # Filter out control characters and other non-XML-compatible characters
            filtered_line = re.sub(r'[^\x20-\x7E\x0A\x0D]', '', line)
            docx.add_paragraph(filtered_line)
    output_file = "output.docx"
    docx.save(output_file)
    return output_file



def perform_ocr_pdf_to_text(file):
    print("Performing OCR on each page in the PDF file...")
    pdf_document = fitz.open(stream=file.read(), filetype="pdf")
    text = ""
    for page_number in range(len(pdf_document)):
        page = pdf_document.load_page(page_number)
        pix = page.get_pixmap()
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        text += pytesseract.image_to_string(img)
    output_file = "output.txt"
    with open(output_file, 'w') as f:
        f.write(text)
    return output_file


def perform_ocr_pdf_to_xlsx(file):
    print("Performing OCR on each page in the PDF file...")
    pdf_document = fitz.open(stream=file.read(), filetype="pdf")
    data = []
    for page_number in range(len(pdf_document)):
        page = pdf_document.load_page(page_number)
        pix = page.get_pixmap()
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        text = pytesseract.image_to_string(img)
        # Filter out illegal characters for Excel
        filtered_text = re.sub(r'[^\x20-\x7E\x0A\x0D]', '', text)
        data.append(filtered_text.split('\n'))
    df = pd.DataFrame(data)
    output_file = "output.xlsx"
    df.to_excel(output_file, index=False, header=False)
    return output_file

def translate_text_cloudflare(text, source_lang, target_lang):
    model = "@cf/meta/m2m100-1.2b"

    account_id = CLOUDFLARE_ACCOUNT_ID
    api_token = CLOUDFLARE_API_TOKEN
    
    print(CLOUDFLARE_API_TOKEN)
    print(CLOUDFLARE_ACCOUNT_ID)
    response = requests.post(
        f"https://api.cloudflare.com/client/v4/accounts/{account_id}/ai/run/{model}",
        headers={"Authorization": f"Bearer {api_token}"},
        json={
            "text": text,
            "source_lang": source_lang,
            "target_lang": target_lang
        }
    )
    print(response.status_code)
    print(response.json())
    
    inference = response.json()
    print(inference["result"]["translated_text"])
    return inference["result"]["translated_text"]