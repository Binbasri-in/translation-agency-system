from django.shortcuts import render, HttpResponse
from .forms import DocumentForm
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import arabic_reshaper
from bidi.algorithm import get_display
import pytesseract
from PIL import Image
import pandas as pd
import re


def upload_file(request):
    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            input_file_type = form.cleaned_data['input_file_type']
            print(f"Input file type: {input_file_type}")
            output_file_type = form.cleaned_data['output_file_type']
            print(f"Output file type: {output_file_type}")
            uploaded_file = request.FILES['docfile']
            print(f"Uploaded file: {uploaded_file.name}")

            # confirm that the file is a PDF or an image
            file_extension = uploaded_file.name.split('.')[-1].lower()
            if file_extension not in ['pdf', 'jpg', 'jpeg', 'png']:
                error_message = "Invalid file format. Only PDFs and images are supported."
                return render(request, 'ocr_app/upload.html', {'form': form, 'error': error_message})

            # Perform OCR on the uploaded file
            if input_file_type == 'pdf' and output_file_type == 'docx':
                output_file = perform_ocr_pdf_to_docx(uploaded_file)
            elif input_file_type == 'pdf' and output_file_type == 'txt':
                output_file = perform_ocr_pdf_to_text(uploaded_file)
            elif input_file_type == 'pdf' and output_file_type == 'xlsx':
                output_file = perform_ocr_pdf_to_xlsx(uploaded_file)
            elif input_file_type == 'image' and output_file_type == 'docx':
                output_file = perform_ocr_image_to_docx(uploaded_file)
            elif input_file_type == 'image' and output_file_type == 'txt':
                output_file = perform_ocr_image_to_text(uploaded_file)
            elif input_file_type == 'image' and output_file_type == 'xlsx':
                output_file = perform_ocr_image_to_xlsx(uploaded_file)
            else:
                error_message = "Invalid combination of input and output file types."
                return render(request, 'ocr_app/upload.html', {'form': form, 'error': error_message})

            print(f"OCR output saved to {output_file}")

            if output_file is None:
                print("Error during OCR process")
                error_message = "Error during OCR process"
                return render(request, 'ocr_app/upload.html', {'form': form, 'error': error_message})

            # Return the text file to the user for download with explicit UTF-8 encoding
            with open(output_file, 'rb') as f:
                response = HttpResponse(f.read(), content_type='text/plain; charset=utf-8')
                response['Content-Disposition'] = f'attachment; filename="{output_file}"'
                print("Returning response")
                return response
        else:
            error_message = "Invalid file format. Only PDFs or Images are supported."
            return render(request, 'ocr_app/upload.html', {'form': form, 'error': error_message})
    else:
        error_message = ""

    form = DocumentForm()

    return render(request, 'ocr_app/upload.html', {'form': form, 'error': error_message})


def perform_ocr_image_to_text(file):
    print("Performing OCR on the image...")
    image = Image.open(file)
    text = pytesseract.image_to_string(image)
    output_file = "output.txt"
    with open(output_file, 'w') as f:
        f.write(text)
    return output_file


def perform_ocr_image_to_docx(file):
    print("Performing OCR on the image...")
    image = Image.open(file)
    text = pytesseract.image_to_string(image)
    # Filter out control characters and other non-XML-compatible characters
    filtered_text = re.sub(r'[^\x20-\x7E\x0A\x0D]', '', text)
    docx = DocxDocument()
    docx.add_paragraph(filtered_text)
    output_file = "output.docx"
    docx.save(output_file)
    return output_file


def perform_ocr_image_to_xlsx(file):
    print("Performing OCR on the image...")
    image = Image.open(file)
    text = pytesseract.image_to_string(image)
    # Filter out illegal characters for Excel
    filtered_text = re.sub(r'[^\x20-\x7E\x0A\x0D]', '', text)
    df = pd.DataFrame([filtered_text.split('\n')])
    output_file = "output.xlsx"
    df.to_excel(output_file, index=False, header=False)
    return output_file



def perform_ocr_pdf_to_docx(file):
    print("Performing OCR on each page in the PDF file...")
    pdf_document = fitz.open(stream=file.read(), filetype="pdf")
    docx = DocxDocument()  # Create a new Word document
    for page_number in range(len(pdf_document)):
        page = pdf_document.load_page(page_number)
        pix = page.get_pixmap()
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        text = pytesseract.image_to_string(img)
        # Reshape and reorder Arabic text
        reshaped_text = arabic_reshaper.reshape(text)
        for line in reshaped_text.split('\n'):
            # Filter out control characters and other non-XML-compatible characters
            filtered_line = re.sub(r'[^\x20-\x7E\x0A\x0D]', '', line)
            docx.add_paragraph(filtered_line)
    output_file = "output.docx"
    docx.save(output_file)
    return output_file



def perform_ocr_pdf_to_text(file):
    print("Performing OCR on each page in the PDF file...")
    pdf_document = fitz.open(stream=file.read(), filetype="pdf")
    text = ""
    for page_number in range(len(pdf_document)):
        page = pdf_document.load_page(page_number)
        pix = page.get_pixmap()
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        text += pytesseract.image_to_string(img)
    output_file = "output.txt"
    with open(output_file, 'w') as f:
        f.write(text)
    return output_file


def perform_ocr_pdf_to_xlsx(file):
    print("Performing OCR on each page in the PDF file...")
    pdf_document = fitz.open(stream=file.read(), filetype="pdf")
    data = []
    for page_number in range(len(pdf_document)):
        page = pdf_document.load_page(page_number)
        pix = page.get_pixmap()
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        text = pytesseract.image_to_string(img)
        # Filter out illegal characters for Excel
        filtered_text = re.sub(r'[^\x20-\x7E\x0A\x0D]', '', text)
        data.append(filtered_text.split('\n'))
    df = pd.DataFrame(data)
    output_file = "output.xlsx"
    df.to_excel(output_file, index=False, header=False)
    return output_file